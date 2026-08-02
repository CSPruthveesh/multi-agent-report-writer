"""Turn a finished report into a .docx or a .pdf.

ONE PARSER, TWO RENDERERS
-------------------------
The markdown the Writer emits is small and known: one H1, H2 sections, paragraphs,
bold, horizontal rules, bullets under "Known limitations", and [F001] citations. So
`blocks()` reads it once into a flat list and each renderer walks that list.

The alternative — a markdown reader inside each exporter — is the mistake the front end
already made and fixed: the trace was rendered by two copies of the same code and they
drifted until the pages disagreed about what a run looked like. A report that paginates
differently depending on which button was pressed is the same defect with a longer
feedback loop.

FONTS
-----
The page lets a reader put the report in any font their machine has, and both formats
keep it.

  .docx  carries the family by name. Word resolves it wherever the file opens, which is
         the right behaviour for a document that travels.
  .pdf   embeds the actual font file. The machine's font directory is indexed once, the
         requested family's regular and bold faces are registered with reportlab, and
         the glyphs travel inside the PDF — so it looks the same on a machine that has
         never heard of Palatino Linotype.

Converting the .docx through Word or LibreOffice would also preserve the font and was
the obvious first idea. Embedding is better on every axis that matters here: no second
program to install, nothing to shell out to, no temp files, no Windows-only COM
automation, and it takes about a tenth of a second rather than several.

A family the server cannot find falls back to the nearest of reportlab's three built-in
families. That is not hypothetical — Archivo and JetBrains Mono are webfonts this page
loads over the network, so the browser can offer them while the operating system has
never installed them. The fallback is a near match by design: a grotesque sans becomes
Helvetica, a monospace becomes Courier.
"""

from __future__ import annotations

import io
import logging
import os
import re
import sys
from functools import lru_cache
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from reportlab.lib.enums import TA_JUSTIFY
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont, TTFontFile
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    SimpleDocTemplate,
    Spacer,
)
from reportlab.platypus import Paragraph as RLParagraph

log = logging.getLogger(__name__)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
CITE_RE = re.compile(r"\[((?:F\d{3})(?:,\s*F\d{3})*)\]")

# What the page offers, grouped by what a PDF can actually do with it. Anything not
# named here falls back to Helvetica, which is also what an unknown font should do.
_SERIF = {
    "georgia", "times new roman", "cambria", "constantia", "palatino linotype",
    "book antiqua", "garamond", "sylfaen", "charter", "serif",
}
_MONO = {"consolas", "courier new", "jetbrains mono", "monospace"}


def pdf_family(font: str | None) -> str:
    """The built-in reportlab family a requested font falls back to."""
    name = (font or "").strip().lower()
    if name in _MONO:
        return "Courier"
    if name in _SERIF:
        return "Times-Roman"
    return "Helvetica"


def _font_dirs() -> list[Path]:
    """Where this operating system keeps its fonts."""
    if sys.platform == "win32":
        win = Path(os.environ.get("WINDIR", r"C:\Windows"))
        local = os.environ.get("LOCALAPPDATA")
        dirs = [win / "Fonts"]
        if local:  # per-user installs, which is where a font added today lands
            dirs.append(Path(local) / "Microsoft" / "Windows" / "Fonts")
        return dirs
    if sys.platform == "darwin":
        return [Path("/System/Library/Fonts"), Path("/Library/Fonts"),
                Path.home() / "Library" / "Fonts"]
    return [Path("/usr/share/fonts"), Path("/usr/local/share/fonts"),
            Path.home() / ".fonts", Path.home() / ".local/share/fonts"]


@lru_cache(maxsize=1)
def font_index() -> dict[str, dict[str, str]]:
    """family (lowercased) -> {"regular": path, "bold": path}.

    Built by reading each file's name table rather than by guessing from filenames,
    because they do not correspond: Book Antiqua lives in BKANT.TTF and Palatino
    Linotype in pala.ttf. reportlab already has the parser, so this needs no new
    dependency.

    Cached for the life of the process. Reading ~320 files costs about 1.3 seconds and
    a font is not installed halfway through a session — but the cost lands on whoever
    exports first, so it is deliberately not done at import time.
    """
    out: dict[str, dict[str, str]] = {}
    for d in _font_dirs():
        if not d.is_dir():
            continue
        for p in sorted(d.glob("*.ttf")):
            try:
                f = TTFontFile(str(p))
            except Exception as e:  # noqa: BLE001
                # A font directory holds bitmap fonts, broken files and formats
                # reportlab does not read. Skipping is right; skipping silently is
                # not — a family missing from the picker's PDF should be findable.
                log.debug("skipping font %s: %s", p.name, e)
                continue
            fam = _text(f.familyName).lower()
            style = _text(f.styleName).lower()
            if "italic" in style or "oblique" in style:
                continue          # nothing here sets italics
            slot = "bold" if "bold" in style else "regular"
            out.setdefault(fam, {}).setdefault(slot, str(p))
    return out


def _text(v: bytes | str) -> str:
    return v.decode("latin-1", "replace") if isinstance(v, bytes) else str(v)


_registered: set[str] = set()


def resolve_pdf_font(font: str | None) -> tuple[str, str, str]:
    """(regular, bold, applied) — embedding the real font where it can be found.

    `applied` is the family the document actually ends up in, which is what the caller
    reports back. It is the requested name when the font was embedded and reportlab's
    substitute when it was not, so the two cases are distinguishable by the answer
    rather than only by the outcome.
    """
    name = (font or "").strip()
    entry = font_index().get(name.lower()) if name else None
    if not entry or "regular" not in entry:
        fam = pdf_family(name)
        bold = "Times-Bold" if fam == "Times-Roman" else f"{fam}-Bold"
        return fam, bold, fam

    reg, bld = f"emb-{name}", f"emb-{name}-bold"
    if reg not in _registered:
        pdfmetrics.registerFont(TTFont(reg, entry["regular"]))
        pdfmetrics.registerFont(TTFont(bld, entry.get("bold", entry["regular"])))
        # So <b> inside a paragraph resolves to the bold face rather than being
        # synthesised — reportlab looks the family up, it does not embolden.
        pdfmetrics.registerFontFamily(reg, normal=reg, bold=bld,
                                      italic=reg, boldItalic=bld)
        _registered.add(reg)
    return reg, bld, name


def blocks(md: str) -> list[tuple[str, str]]:
    """Flatten the report into (kind, text) pairs.

    Kinds: h1, h2, rule, li, p. Line endings are normalised first — the same \\r\\n
    assumption that collapsed the whole report into one block in the browser renderer
    would collapse it into one paragraph here.
    """
    out: list[tuple[str, str]] = []
    for chunk in re.split(r"\n{2,}", md.replace("\r\n", "\n").replace("\r", "\n")):
        chunk = chunk.strip()
        if not chunk:
            continue
        if chunk.startswith("# "):
            out.append(("h1", chunk[2:].strip()))
        elif chunk.startswith("## "):
            out.append(("h2", chunk[3:].strip()))
        elif re.fullmatch(r"-{3,}", chunk):
            out.append(("rule", ""))
        elif re.match(r"^\s*[-*]\s", chunk):
            # A wrapped bullet continues the item above it rather than opening an
            # empty one — the same rule the browser renderer applies.
            for line in chunk.split("\n"):
                m = re.match(r"^\s*[-*]\s+(.*)$", line)
                if m:
                    out.append(("li", m.group(1).strip()))
                elif out and out[-1][0] == "li" and line.strip():
                    out[-1] = ("li", f"{out[-1][1]} {line.strip()}")
        else:
            out.append(("p", " ".join(chunk.split())))
    return out


def _runs(text: str) -> list[tuple[str, bool]]:
    """Split on **bold** into (text, is_bold) pairs."""
    parts, last = [], 0
    for m in BOLD_RE.finditer(text):
        if m.start() > last:
            parts.append((text[last:m.start()], False))
        parts.append((m.group(1), True))
        last = m.end()
    if last < len(text):
        parts.append((text[last:], False))
    return parts or [(text, False)]


# ------------------------------------------------------------------------ docx
def to_docx(md: str, *, title: str = "", font: str | None = None) -> bytes:
    doc = Document()
    if font:
        # The Normal style, so every paragraph inherits it rather than each one
        # carrying its own run-level override.
        doc.styles["Normal"].font.name = font
    doc.styles["Normal"].font.size = Pt(11)

    if title:
        doc.core_properties.title = title

    for kind, text in blocks(md):
        if kind == "rule":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("* * *")
            continue
        if kind == "h1":
            doc.add_heading(text, level=0)
            continue
        if kind == "h2":
            doc.add_heading(text, level=1)
            continue

        p = doc.add_paragraph(style="List Bullet" if kind == "li" else None)
        # Citations get their own runs so they can be coloured. They are the reason to
        # trust the document, and in a printed report they are the only thing that
        # links a claim back to a source.
        for chunk, bold in _runs(text):
            for i, piece in enumerate(CITE_RE.split(chunk)):
                if not piece:
                    continue
                run = p.add_run(f"[{piece}]" if i % 2 else piece)
                run.bold = bold
                if i % 2:
                    run.font.color.rgb = RGBColor(0x0F, 0x6E, 0x6B)
                    run.font.size = Pt(9)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ------------------------------------------------------------------------- pdf
def _xml(text: str) -> str:
    """Escape for reportlab's mini-HTML, then put the bold tags back."""
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return BOLD_RE.sub(r"<b>\1</b>", safe)


def to_pdf(md: str, *, title: str = "", font: str | None = None) -> bytes:
    fam, bold, _ = resolve_pdf_font(font)
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "body", parent=base["BodyText"], fontName=fam, fontSize=10.5, leading=15.5,
        alignment=TA_JUSTIFY, spaceAfter=9,
    )
    h1 = ParagraphStyle(
        "h1", parent=base["Title"], fontName=bold,
        fontSize=19, leading=23, spaceAfter=14, alignment=0,
    )
    h2 = ParagraphStyle(
        "h2", parent=base["Heading2"], fontName=bold,
        fontSize=12.5, leading=16, spaceBefore=13, spaceAfter=5,
    )

    flow: list[Any] = []
    bullets: list[ListItem] = []

    def flush() -> None:
        if bullets:
            # bulletFontName as well: it defaults to Helvetica whatever the body is set
            # in, so the dots beside a Times list were drawn in a different face from
            # the text beside them.
            flow.append(ListFlowable(list(bullets), bulletType="bullet",
                                     leftIndent=14, bulletFontSize=7,
                                     bulletFontName=fam))
            flow.append(Spacer(1, 7))
            bullets.clear()

    for kind, text in blocks(md):
        if kind != "li":
            flush()
        if kind == "h1":
            flow.append(RLParagraph(_xml(text), h1))
        elif kind == "h2":
            flow.append(RLParagraph(_xml(text), h2))
        elif kind == "rule":
            flow.append(Spacer(1, 5))
            flow.append(HRFlowable(width="100%", thickness=0.6, color="#C3CACC"))
            flow.append(Spacer(1, 9))
        elif kind == "li":
            bullets.append(ListItem(RLParagraph(_xml(text), body), leftIndent=14))
        else:
            flow.append(RLParagraph(_xml(text), body))
    flush()

    buf = io.BytesIO()
    SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=24 * mm, rightMargin=24 * mm,
        topMargin=20 * mm, bottomMargin=20 * mm,
        title=title or "Report", author="Multi-Agent Report Writer",
    ).build(flow)
    return buf.getvalue()
